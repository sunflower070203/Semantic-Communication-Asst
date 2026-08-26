#!/usr/bin/env python3
"""Convert the project plan markdown to DOCX and PDF (submission format).

Supports the subset used in docs/plan_book.md:
  # / ## / ### headings, --- rules, | tables |, N. lists, **bold**, code blocks.

Usage: python md_to_docx_pdf.py <input.md> <output_dir>
"""

import os
import re
import sys


def parse_lines(text):
    """Yield (kind, content) tuples for the supported markdown subset."""
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue
        if line.strip() == "---":
            yield ("rule", "")
            i += 1
            continue
        if line.lstrip().startswith("```"):
            buf = []
            i += 1
            while i < len(lines) and not lines[i].lstrip().startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1
            yield ("code", "\n".join(buf))
            continue
        if line.startswith("|") and "|" in line[1:]:
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(lines[i].strip())
                i += 1
            yield ("table", rows)
            continue
        m = re.match(r"^(#{1,3})\s+(.*)$", line)
        if m:
            yield ("h" + str(len(m.group(1))), m.group(2))
            i += 1
            continue
        m = re.match(r"^(\d+)\.\s+(.*)$", line)
        if m:
            yield ("ol", m.group(2))
            i += 1
            continue
        m = re.match(r"^[-*]\s+(.*)$", line)
        if m:
            yield ("ul", m.group(1))
            i += 1
            continue
        yield ("p", line)
        i += 1


def split_bold(text):
    """Split text into (bold: bool, chunk) segments."""
    parts = re.split(r"(\*\*.+?\*\*)", text)
    out = []
    for p in parts:
        if not p:
            continue
        if p.startswith("**") and p.endswith("**") and len(p) > 4:
            out.append((True, p[2:-2]))
        else:
            out.append((False, p))
    return out


def build_docx(items, out_path):
    import docx
    from docx.shared import Pt, RGBColor

    doc = docx.Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)
    hcolors = {
        "h1": RGBColor(0x1F, 0x3A, 0x5F),
        "h2": RGBColor(0x1F, 0x3A, 0x5F),
        "h3": RGBColor(0x2E, 0x5E, 0x8C),
    }
    for kind, content in items:
        if kind == "h1":
            p = doc.add_heading(level=1)
            r = p.add_run(content)
            r.font.size = Pt(20)
            r.font.color.rgb = hcolors["h1"]
        elif kind == "h2":
            p = doc.add_heading(level=2)
            r = p.add_run(content)
            r.font.size = Pt(15)
            r.font.color.rgb = hcolors["h2"]
        elif kind == "h3":
            p = doc.add_heading(level=3)
            r = p.add_run(content)
            r.font.size = Pt(12.5)
            r.font.color.rgb = hcolors["h3"]
        elif kind == "p":
            p = doc.add_paragraph()
            for bold, chunk in split_bold(content):
                r = p.add_run(chunk)
                r.bold = bold
        elif kind in ("ol", "ul"):
            p = doc.add_paragraph(style="List Number" if kind == "ol" else "List Bullet")
            for bold, chunk in split_bold(content):
                r = p.add_run(chunk)
                r.bold = bold
        elif kind == "code":
            p = doc.add_paragraph()
            r = p.add_run(content)
            r.font.name = "Consolas"
            r.font.size = Pt(9)
        elif kind == "table":
            rows = []
            for row in content:
                cells = [c.strip() for c in row.strip("|").split("|")]
                rows.append(cells)
            if len(rows) < 2:
                continue
            header = rows[0]
            body = [r for r in rows[1:] if not set(r) == {"---"} and not all(c == "---" for c in r)]
            table = doc.add_table(rows=1 + len(body), cols=len(header))
            table.style = "Light Grid Accent 1"
            for j, cell in enumerate(header):
                for bold, chunk in split_bold(cell):
                    run = table.rows[0].cells[j].paragraphs[0].add_run(chunk)
                    run.bold = True
            for i, row in enumerate(body):
                for j, cell in enumerate(row):
                    if j < len(header):
                        table.rows[i + 1].cells[j].text = re.sub(r"\*\*", "", cell)
        elif kind == "rule":
            doc.add_paragraph()
    doc.save(out_path)


def build_pdf(items, out_path):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.platypus import (
        SimpleDocTemplate,
        Paragraph,
        Spacer,
        Table,
        TableStyle,
        Preformatted,
    )

    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    base = dict(fontName="STSong-Light", wordWrap="CJK")
    s_title = ParagraphStyle("t", **base, fontSize=22, leading=30, textColor=colors.HexColor("#1F3A5F"))
    s_h2 = ParagraphStyle("h2", **base, fontSize=15, leading=22, spaceBefore=12, spaceAfter=6, textColor=colors.HexColor("#1F3A5F"))
    s_h3 = ParagraphStyle("h3", **base, fontSize=12.5, leading=19, spaceBefore=8, spaceAfter=4, textColor=colors.HexColor("#2E5E8C"))
    s_body = ParagraphStyle("body", **base, fontSize=10.5, leading=17, spaceAfter=4)
    s_code = ParagraphStyle("code", fontName="Courier", fontSize=8.5, leading=12, backColor=colors.HexColor("#F4F4F4"), borderPadding=6)

    doc = SimpleDocTemplate(out_path, pagesize=A4, leftMargin=22 * mm, rightMargin=22 * mm, topMargin=20 * mm, bottomMargin=20 * mm)
    story = []
    for kind, content in items:
        if kind == "h1":
            story.append(Paragraph(content, s_title))
            story.append(Spacer(1, 8))
        elif kind == "h2":
            story.append(Paragraph(content, s_h2))
        elif kind == "h3":
            story.append(Paragraph(content, s_h3))
        elif kind == "p":
            html = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", content)
            story.append(Paragraph(html, s_body))
        elif kind in ("ol", "ul"):
            html = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", content)
            story.append(Paragraph(("• " if kind == "ul" else "• ") + html, s_body))
        elif kind == "code":
            story.append(Preformatted(content, s_code))
            story.append(Spacer(1, 4))
        elif kind == "table":
            rows = []
            for row in content:
                rows.append([re.sub(r"\*\*", "", c.strip()) for c in row.strip("|").split("|")])
            rows = [r for r in rows if not all(c == "---" for c in r)]
            if len(rows) < 2:
                continue
            widths = None
            tbl = Table(rows)
            tbl.setStyle(
                TableStyle(
                    [
                        ("FONTNAME", (0, 0), (-1, 0), "STSong-Light"),
                        ("FONTNAME", (0, 1), (-1, -1), "STSong-Light"),
                        ("FONTSIZE", (0, 0), (-1, -1), 9),
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF5")),
                        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#B8C4D0")),
                        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 5),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                    ]
                )
            )
            story.append(tbl)
            story.append(Spacer(1, 6))
        elif kind == "rule":
            story.append(Spacer(1, 4))
    doc.build(story)


def main():
    src = sys.argv[1]
    out_dir = sys.argv[2]
    os.makedirs(out_dir, exist_ok=True)
    text = open(src, encoding="utf-8").read()
    items = list(parse_lines(text))
    base = os.path.splitext(os.path.basename(src))[0]
    docx_path = os.path.join(out_dir, base + ".docx")
    pdf_path = os.path.join(out_dir, base + ".pdf")
    build_docx(items, docx_path)
    build_pdf(items, pdf_path)
    print("DOCX:", docx_path)
    print("PDF:", pdf_path)


if __name__ == "__main__":
    main()
